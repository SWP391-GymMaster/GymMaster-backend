"""Dien Template5_Final Release Document.

Muc I liet ke goi ban giao; II Installation Guides; III User Manual.
Ten file theo quy uoc thay: XYZ = ma nhom + ma du an (vd G1-GYM).

Dung: python fill_final_release.py <template.docx> <out.docx> [--code GYM]
"""
import argparse
import shutil
import subprocess
import sys
from datetime import date

try:
    import docx
    from docx.shared import Pt
except ImportError:
    sys.exit('Thieu python-docx. Chay: uv run --with python-docx python ' + __file__)

BE = r'D:\GymMaster\GymMaster-backend'
FE = r'D:\GymMaster\GymMaster-frontend'


def sha(repo):
    return subprocess.run(['git', 'rev-parse', '--short', 'HEAD'], cwd=repo,
                          capture_output=True, text=True).stdout.strip()


def set_cell(cell, text):
    cell.text = ''
    p = cell.paragraphs[0]
    for i, line in enumerate(str(text).split('\n')):
        r = p.add_run(line)
        r.font.size = Pt(10)
        if i < len(str(text).split('\n')) - 1:
            r.add_break()


def uniq(row):
    out, seen = [], set()
    for c in row.cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            out.append(c)
    return out


def fill(t, rows):
    while len(t.rows) - 1 < len(rows):
        t.add_row()
    while len(t.rows) - 1 > len(rows):
        t._tbl.remove(t.rows[-1]._tr)
    for i, vals in enumerate(rows, 1):
        cs = uniq(t.rows[i])
        for j, v in enumerate(vals):
            if j < len(cs):
                set_cell(cs[j], v)


def para_after(d, marker, texts):
    """Chen cac doan ngay sau doan chua marker."""
    for p in d.paragraphs:
        if marker in p.text:
            anchor = p._p
            for style, txt in texts:
                q = d.add_paragraph(style=style)
                q.add_run(txt)
                anchor.addnext(q._p)
                anchor = q._p
            return True
    return False


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('template')
    ap.add_argument('out')
    ap.add_argument('--code', default='GYM')
    a = ap.parse_args()

    X = a.code
    shutil.copy(a.template, a.out)
    d = docx.Document(a.out)

    # --- I. Deliverable Package
    fill(d.tables[0], [
        ['1', '{}_DB_final.sql'.format(X),
         'Script SQL Server: tao 23 bang + du lieu. Nguon: database/*.sql '
         '(011 script tang dan) + GymMaster_SQLServer_Final.sql'],
        ['2', '{}_RDS_final.docx'.format(X),
         'Requirement & Design Specification — 3 phan: Overview (5 actor, 29 use case, '
         '46 man, phan quyen, 23 bang DB, 15 package), Requirement Spec (29 bang UC + '
         'business rules), Design Spec (46 man: mockup + bang field + Database Access)'],
        ['3', '{}_SDS_final.docx'.format(X),
         'Software Design Specification — package diagram, database design, '
         'class specification cho 10 feature backend'],
        ['4', '{}_Final Product Backlog.xlsx'.format(X),
         'Project Tracking: 59 dong (46 man + 3 non-UI function + 10 API backend). '
         'Sheet Project + Iter1-4. Cot Actual lay tu ngay commit that.'],
        ['5', '{}_Issues Report.xlsx'.format(X),
         'Issues Report: 215 dong sinh tu git log that (co link commit GitHub kiem chung duoc), '
         'phan theo milestone iter1-4, nhan Task/Defect.'],
        ['6', '{}_UseCase.drawio'.format(X),
         'Use case diagram — 5 tab (Admin/Staff/PT/Member/System), mo bang draw.io'],
    ])

    tag_be, tag_fe = sha(BE), sha(FE)
    for p in d.paragraphs:
        if 'git lab tag link' in p.text or 'gitlab tag' in p.text.lower():
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = (
                    'Tagged source codes:\n'
                    '  Backend : https://github.com/SWP391-GymMaster/GymMaster-backend '
                    '(commit {})\n'
                    '  Frontend: https://github.com/SWP391-GymMaster/GymMaster-frontend '
                    '(commit {})\n'
                    '  [CAN BO SUNG] Guide yeu cau tag qua GitLab — du an dung GitHub, '
                    'can hoi thay.'.format(tag_be, tag_fe))
        if 'YouTube link' in p.text:
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = 'Demonstration video: [CAN BO SUNG] chua quay'

    # --- II. Installation Guides
    para_after(d, 'II. Installation Guides', [
        (None, 'Yeu cau: .NET SDK 10 · Node.js 24 · SQL Server. '
               '(Guide cua thay ghi NetBeans 8.2 — khong ap dung, du an la .NET + Next.js.)'),
        (None, '1. Database: chay cac script trong database/ theo thu tu so (004 -> 013), '
               'hoac dung GymMaster_SQLServer_Final.sql de tao moi.'),
        (None, '2. Backend: dat secret bang User Secrets (KHONG commit):'),
        (None, '   dotnet user-secrets set "ConnectionStrings:DefaultConnection" "..." '
               '--project backend/GymMaster.API'),
        (None, '   Cac key khac: Jwt:Key, VnPay:TmnCode/HashSecret, Gemini:ApiKey, '
               'Cloudinary:*, Email:* (SMTP App Password).'),
        (None, '   Chay: cd backend/GymMaster.API && dotnet run   ->  http://localhost:5042'),
        (None, '   OpenAPI: http://localhost:5042/openapi/v1.json (.NET 10 native, KHONG co /swagger)'),
        (None, '3. Frontend: tao .env.local voi NEXT_PUBLIC_API_BASE_URL=http://localhost:5042 '
               'va NEXT_PUBLIC_API_MOCKING= (de trong = goi backend that).'),
        (None, '   npm ci && npm run dev   ->  http://localhost:3000/login'),
        (None, '4. Tai khoan demo (seeder tu tao): admin@gymmaster.local/Admin123! · '
               'staff@gymmaster.local/Staff123! · pt@gymmaster.local/Pt123! · '
               'member@gymmaster.local/Member123!'),
        (None, '5. Deploy: Google Cloud Run + Cloud SQL — xem docs/05-Deployment/deploy-gcp.md.'),
    ])

    # --- III. User Manual
    para_after(d, '1. Overview', [
        (None, 'GymMaster la he thong web quan ly vong doi hoi vien phong gym (1 chi nhanh, '
               '~1000 hoi vien). 4 vai tro: Admin, Staff, PT, Member.'),
        (None, 'Luong nghiep vu chinh: Admin/Staff tao ho so hoi vien -> ban hoac gia han goi tap '
               '-> ghi nhan thanh toan (tien mat hoac VNPay) -> membership Active -> hoi vien '
               'check-in -> Admin phan cong PT -> PT tao giao an va ghi chu -> hoi vien xem '
               'tien do va ghi nhat ky bua an -> Admin xem dashboard va audit log.'),
        ('Heading 2', '2. Luong ban va kich hoat goi tap'),
        (None, '[CAN BO SUNG] Chen anh chup man hinh tung buoc. Anh co san o '
               '../GymMaster-frontend/docs/design/screenshots/ (43 anh): staff-sell-package.png, '
               'staff-payments.png, member-membership.png, admin-memberships.png.'),
        ('Heading 2', '3. Luong check-in'),
        (None, '[CAN BO SUNG] Anh: staff-check-in.png, pt-check-in.png. '
               'Gioi han 2 luot/ngay (DAILY_LIMIT_REACHED 409), moc ngay theo GMT+7.'),
        ('Heading 2', '4. Luong PT training'),
        (None, '[CAN BO SUNG] Anh: admin-assignments.png, pt-workout-planner.png, '
               'pt-trainer-notes.png, member-workout.png.'),
        ('Heading 2', '5. Luong dinh duong + quet anh AI'),
        (None, '[CAN BO SUNG] Anh: member-meal-journal.png, member-nutrition-summary.png. '
               'Quet anh dung Gemini Vision gemini-2.5-flash.'),
    ])

    for p in d.paragraphs[:6]:
        if 'Hanoi' in p.text and '202' in p.text:
            for r in p.runs[1:]:
                r.text = ''
            if p.runs:
                p.runs[0].text = '– Hanoi, {} –'.format(date.today().strftime('%B %Y'))

    d.save(a.out)
    print('Da dien Final Release -> {}'.format(a.out))
    print('  commit backend : {}'.format(tag_be))
    print('  commit frontend: {}'.format(tag_fe))
    print('  CAN BO SUNG    : link tag GitLab (du an dung GitHub), video demo, anh User Manual')


if __name__ == '__main__':
    main()
