"""Dien RDS phan III (Design Specifications) — XOA MAU cua thay roi thay bang
noi dung that cua GymMaster.

Nguon:
  Mockup UI      <- ../GymMaster-frontend/docs/design/screenshots/*.png (43 anh)
  Bang field     <- Zod schema trong ../GymMaster-frontend/src/features/*/schemas/
  Database Access<- _dbContext.<Bang> trong Features/<Feature>/*Service.cs

CANH BAO: template co san muc mau cua du an GAMS (User Login, Setting List,
Setting Details, anh co chu "Teacher"). Khong xoa la nop nham du an nguoi khac.

Dung: python fill_rds_design.py <rds.docx> <inventory.csv> [--fe <path>]
"""
import argparse
import copy
import csv
import re
import sys
from pathlib import Path

try:
    import docx
    from docx.oxml.ns import qn
    from docx.shared import Pt, Inches
except ImportError:
    sys.exit('Thieu python-docx. Chay: uv run --with python-docx python ' + __file__)

sys.path.insert(0, str(Path(__file__).parent))
from fill_tracking import NAMES
from who_owns import FE_FEATURE, pick

API = Path('backend/GymMaster.API/Features')

# ten anh <- route (khop cach dat ten trong visual-screenshots.spec.ts)
SHOT = {
    '/': 'landing', '/ (landing)': 'landing', '/about': 'about', '/login': 'login',
    '/signup': 'signup', '/forgot-password': 'forgot-password',
    '/reset-password': 'reset-password', '/welcome': 'welcome',
    '/change-password': 'change-password',
    # Route dong: chup o id 101 (member dau tien trong mock data).
    '/admin/members/[id]': 'admin-member-360',
    '/staff/members/[id]': 'staff-member-360',
    '/member/membership/vnpay-return': 'member-vnpay-return',
    '/admin/dashboard': 'admin-dashboard', '/admin/users': 'admin-users',
    '/admin/staff': 'admin-staff', '/admin/trainers': 'admin-trainers',
    '/admin/members': 'admin-members', '/admin/assignments': 'admin-assignments',
    '/admin/audit-logs': 'admin-audit-logs', '/admin/packages': 'admin-packages',
    '/admin/memberships': 'admin-memberships', '/admin/payments': 'admin-payments',
    '/admin/notifications': 'admin-notifications', '/admin/profile': 'admin-profile',
    '/staff/dashboard': 'staff-dashboard', '/staff/members': 'staff-members',
    '/staff/check-in': 'staff-check-in', '/staff/sell-package': 'staff-sell-package',
    '/staff/renew-package': 'staff-renew-package', '/staff/payments': 'staff-payments',
    '/staff/profile': 'staff-profile',
    '/pt/dashboard': 'pt-dashboard', '/pt/members': 'pt-members',
    '/pt/members/[id]': 'pt-member-360', '/pt/members/[id]/workout': 'pt-workout-planner',
    '/pt/members/[id]/notes': 'pt-trainer-notes',
    '/pt/members/[id]/progress': 'pt-member-progress',
    '/pt/check-in': 'pt-check-in', '/pt/profile': 'pt-profile',
    '/member/dashboard': 'member-dashboard', '/member/workout': 'member-workout',
    '/member/nutrition/meal-journal': 'member-meal-journal',
    '/member/nutrition/summary': 'member-nutrition-summary',
    '/member/progress': 'member-progress', '/member/membership': 'member-membership',
    '/member/notes': 'member-notes', '/member/profile': 'member-profile',
    '/member/profile/edit': 'member-profile-edit',
}

# route -> feature backend (de tra bang DB)
FEAT = {
    'login': 'Auth', 'signup': 'Auth', 'forgot-password': 'Auth',
    'reset-password': 'Auth', 'change-password': 'Auth',
    'users': 'Users', 'staff': 'Users', 'members': 'Members', 'trainers': 'Trainers',
    'packages': 'Billing', 'memberships': 'Billing', 'payments': 'Billing',
    'membership': 'Billing', 'sell-package': 'Billing', 'renew-package': 'Billing',
    'check-in': 'CheckIns', 'assignments': 'Training', 'workout': 'Training',
    'notes': 'Training', 'progress': 'Training',
    'nutrition': 'Nutrition', 'meal-journal': 'Nutrition', 'summary': 'Nutrition',
    'dashboard': 'Dashboard', 'audit-logs': 'Dashboard', 'notifications': 'Dashboard',
    'profile': 'Account',
}

# Mot feature co nhieu schema; lay het thi sai. Vd feature 'auth' co ca
# loginSchema lan signupSchema -> man Login bi gan them fullName/phone (field cua
# Signup). Ten schema noi ro no thuoc man nao, nen map thang route -> schema.
# Route khong co o day thi lay toan bo schema cua feature (thua con hon thieu —
# cung nguyen tac voi Business Rules: thua thi nhin thay ma cat).
SCHEMA_OF = {
    '/login': ['loginSchema'],
    '/signup': ['signupSchema'],
    '/forgot-password': ['forgotPasswordSchema'],
    '/reset-password': ['resetPasswordSchema'],
    '/change-password': ['changePasswordSchema'],
    '/admin/packages': ['packageFormSchema'],
    '/admin/users': ['createUserSchema', 'updateUserSchema'],
    '/admin/staff': ['createUserSchema', 'updateUserSchema'],
    '/admin/members': ['createMemberSchema', 'updateMemberSchema', 'memberSearchSchema'],
    '/admin/trainers': ['createTrainerSchema', 'updateTrainerSchema'],
    '/staff/members': ['staffSearchSchema'],
    '/staff/check-in': ['checkInSchema'],
    '/staff/sell-package': ['sellPackageSchema'],
    '/staff/payments': ['manualPaymentSchema'],
    '/pt/check-in': ['checkInSchema'],
    '/pt/members/[id]/notes': ['trainerNoteSchema'],
    '/pt/members/[id]/workout': ['workoutPlanSchema', 'workoutExerciseSchema'],
    '/pt/members/[id]/progress': ['progressEntrySchema'],
    '/member/nutrition/meal-journal': ['mealLogSchema', 'customFoodSchema'],
    '/member/progress': ['progressEntrySchema'],
    '/member/profile/edit': ['memberProfileSchema'],
}

ZOD_FIELD = re.compile(r'^\s*(\w+)\s*:\s*z\s*\.(\w+)\(\s*(?:"([^"]*)")?', re.M)

# Field tro toi schema dung chung thay vi goi z.* truc tiep:
#   fullName: personFieldSchemas.fullName
#   newPassword: passwordSchema
# Khong bat thi accountSchema/changePasswordSchema ra RONG.
ZOD_REF = re.compile(r'^\s*(\w+)\s*:\s*([A-Za-z_][\w.]*)\s*,\s*$', re.M)

# 'z' co the xuong dong truoc '.object' (khi co .refine phia sau):
#   export const resetPasswordSchema = z
#     .object({
# Ban cu doi 'z.object({' lien nhau -> truot resetPasswordSchema va
# changePasswordSchema, hai man do ra rong ma khong ai biet vi sao.
ZOD_OBJ = re.compile(r'export const (\w+Schema)\s*=\s*z\s*\.object\(\{(.*?)^\s*\}\)',
                     re.M | re.S)


def feature_of(route):
    for k, v in FEAT.items():
        if k in route:
            return v
    return None


def tables_of(feature):
    d = API / feature
    if not d.is_dir():
        return []
    seen = set()
    for f in d.glob('*Service.cs'):
        src = f.read_text(encoding='utf-8-sig')
        seen |= set(re.findall(r'_dbContext\.(\w+)\s*\.', src))
    return sorted(seen)


def crud_of(feature):
    d = API / feature
    ops = set()
    for f in d.glob('*Service.cs') if d.is_dir() else []:
        s = f.read_text(encoding='utf-8-sig')
        if re.search(r'\.Add\(|\.AddAsync\(', s):
            ops.add('C')
        if re.search(r'\.(First|Single|Any|Where|ToList|Count|Select)', s):
            ops.add('R')
        if re.search(r'\.Update\(|SaveChangesAsync', s):
            ops.add('U')
        if re.search(r'\.Remove\(|\.RemoveRange\(', s):
            ops.add('D')
    return ''.join(o for o in 'CRUD' if o in ops) or 'R'


def zod_fields(fe_root, route):
    """Field cua man hinh <- Zod schema. Thong bao loi tieng Viet lam mo ta.

    Dung ban do route->feature cua who_owns.py, KHONG doan bang cach so chuoi.
    Ban cu: key = ten thu muc bo dau '-', khop neu key nam trong route da bo '/'.
    Cach do truot phan lon: feature 'billing' vs route '/admin/packages' ->
    'billing' not in 'adminpackages' -> tra ve rong, du billing CO schema.
    """
    want = SCHEMA_OF.get(route)
    if want:
        # Ten schema la duy nhat trong ca code base -> tim khap src/features,
        # KHONG buoc vao thu muc feature. Ban do route->feature khong phai luc nao
        # cung trung cho dat schema: /staff/sell-package map sang feature 'billing'
        # nhung sellPackageSchema lai nam o 'staff-front-desk' -> tim theo feature
        # thi truot.
        search_root = Path(fe_root, 'src/features')
    else:
        feat = pick(route, FE_FEATURE)
        search_root = Path(fe_root, 'src/features', feat) if feat else None
    if not search_root or not search_root.is_dir():
        return []
    out = []
    for sf in search_root.rglob('*.ts'):
        if 'schema' not in sf.name:
            continue
        src = sf.read_text(encoding='utf-8')
        for m in ZOD_OBJ.finditer(src):
            if want and m.group(1) not in want:
                continue
            body = m.group(2)
            for f in ZOD_FIELD.finditer(body):
                name, typ, msg = f.group(1), f.group(2), f.group(3) or ''
                out.append((name, {'string': 'Text Box', 'email': 'Text Box (email)',
                                   'number': 'Number Box', 'boolean': 'Checkbox',
                                   'enum': 'Combo Box', 'date': 'Date Picker',
                                   'coerce': 'Number Box'}.get(typ, typ),
                            msg or 'Truong {} cua form'.format(name)))
            for f in ZOD_REF.finditer(body):
                name, ref = f.group(1), f.group(2)
                low = name.lower()
                typ = ('Date Picker' if 'date' in low else
                       'Combo Box' if low in ('gender', 'status', 'role') else 'Text Box')
                out.append((name, typ,
                            'Truong {} cua form (dung schema chung {})'.format(name, ref)))
    seen, uniq = set(), []
    for f in out:
        if f[0] not in seen:
            seen.add(f[0])
            uniq.append(f)
    return uniq[:14]


def set_cell(cell, text):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.font.size = Pt(9)


def reset_rows(tb, n):
    """Xoa het hang du lieu cua bang mau roi them n hang moi.

    Bang mau cua thay co hang du lieu bi GOP O (vd hang dau bang field chi co 1
    <w:tc> thay vi 3). Giu lai hang do roi ghi de thi cot 2-3 bi nuot mat: ghi
    3 gia tri vao 1 o thi chi gia tri dau tien song. add_row() sinh hang tu luoi
    bang nen luon du cot va khong gop.
    """
    while len(tb.rows) > 1:
        tb._tbl.remove(tb.rows[-1]._tr)
    for _ in range(n):
        tb.add_row()


def uniq_cells(row):
    out, seen = [], set()
    for c in row.cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            out.append(c)
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('rds')
    ap.add_argument('inventory')
    ap.add_argument('--fe', default='../GymMaster-frontend')
    a = ap.parse_args()

    shots = Path(a.fe) / 'docs/design/screenshots'
    d = docx.Document(a.rds)
    body = d.element.body
    kids = list(body.iterchildren())

    # --- 1. XOA MAU: tu sau 'III. Design Specifications' den truoc 'IV. Appendix'
    i_start = i_end = None
    for i, c in enumerate(kids):
        if c.tag == qn('w:p'):
            t = ''.join(n.text or '' for n in c.iter(qn('w:t'))).strip()
            if t.startswith('III. Design Specifications'):
                i_start = i
            elif t.startswith('IV. Appendix') and i_start is not None:
                i_end = i
                break
    if i_start is None or i_end is None:
        sys.exit('Khong tim thay moc phan III / IV')

    # --- 2. LAY BANG MAU TRUOC KHI XOA
    # Tim theo NOI DUNG TIEU DE, khong dung chi so cung.
    # Ban cu ghim d.tables[74]/[75]: so bang phu thuoc so use case cua phan II,
    # nen chi can them/bot 1 use case la index truot -> tpl = None -> khong sinh
    # bang nao, ma van in "46/46" (bien dem no_field khong tang khi tpl la None).
    # Loi im lang: RDS mat sach bang phan III ma script bao thanh cong.
    def _hdr(tbl):
        return [c.text.strip().lower() for c in tbl.rows[0].cells] if tbl.rows else []

    tpl_field = tpl_db = None
    for c in kids[i_start + 1:i_end]:
        if c.tag != qn('w:tbl'):
            continue
        t = docx.table.Table(c, d)
        h = _hdr(t)
        if tpl_field is None and 'field name' in h[0]:
            tpl_field = copy.deepcopy(t._tbl)
        elif tpl_db is None and h and h[0] == 'table':
            tpl_db = copy.deepcopy(t._tbl)
        if tpl_field is not None and tpl_db is not None:
            break

    if tpl_field is None or tpl_db is None:
        sys.exit('Khong tim thay bang mau trong vung III cua template '
                 '(can "Field Name|Field Type|Description" va "Table|CRUD|Description"). '
                 'Template co thay doi?')

    # --- 3. XOA MAU
    removed = 0
    for c in kids[i_start + 1:i_end]:
        body.remove(c)
        removed += 1
    print('Da xoa {} phan tu mau cua thay (User Login / Setting List / Setting '
          'Details + anh "Teacher")'.format(removed))

    with open(a.inventory, encoding='utf-8-sig') as fh:
        inv = [r for r in csv.DictReader(fh) if r['kind'] == 'Screen']

    anchor = kids[i_end]           # chen truoc 'IV. Appendix'
    made = no_shot = no_field = 0

    for n, r in enumerate(inv, 1):
        route = r['name']
        name = NAMES.get(route, route)
        feat = feature_of(route)

        def add(p_el):
            anchor.addprevious(p_el)

        h = d.add_paragraph(style='Heading 3')
        h.add_run('III.{} {}'.format(n, name))
        add(h._p)

        h = d.add_paragraph(style='Heading 5')
        h.add_run('UI Design')
        add(h._p)

        png = shots / (SHOT.get(route, '') + '.png') if SHOT.get(route) else None
        p = d.add_paragraph()
        if png and png.exists():
            p.add_run().add_picture(str(png), width=Inches(5.6))
        else:
            p.add_run('[CAN BO SUNG] Chua co anh chup cho man {}'.format(route))
            no_shot += 1
        add(p._p)

        flds = zod_fields(a.fe, route)
        t = copy.deepcopy(tpl_field)
        add(t)
        tb = docx.table.Table(t, d)
        # Khong co field KHONG phai la thieu sot: 9 man (dashboard, landing,
        # about, welcome, PT Assignment, Member 360) chi HIEN THI du lieu, khong
        # co form nhap lieu nen khong co Zod schema. Ghi '[CAN BO SUNG]' o day la
        # noi sai — nguoi doc se tuong con viec chua lam.
        rows = flds or [('(khong co)', '-',
                         'Man chi hien thi du lieu, khong co form nhap lieu '
                         '(khong co Zod schema trong code)')]
        if not flds:
            no_field += 1
        reset_rows(tb, len(rows))
        for k, vals in enumerate(rows, 1):
            cs = uniq_cells(tb.rows[k])
            for j, v in enumerate(vals):
                if j < len(cs):
                    set_cell(cs[j], v)

        h = d.add_paragraph(style='Heading 5')
        h.add_run('Database Access')
        add(h._p)

        tbls = tables_of(feat) if feat else []
        t = copy.deepcopy(tpl_db)
        add(t)
        tb = docx.table.Table(t, d)
        crud = crud_of(feat) if feat else 'R'
        rows = [(x, crud, 'Truy cap qua {}Service (Features/{})'.format(feat, feat))
                for x in tbls] or [('[CAN BO SUNG]', '-', 'Chua map duoc feature backend')]
        reset_rows(tb, len(rows))
        for k, vals in enumerate(rows, 1):
            cs = uniq_cells(tb.rows[k])
            for j, v in enumerate(vals):
                if j < len(cs):
                    set_cell(cs[j], v)

        p = d.add_paragraph()
        p.add_run('SQL Commands: ').bold = True
        p.add_run('[CAN BO SUNG] Backend dung EF Core LINQ, khong co SQL tho. '
                  'Lay SQL that do EF sinh ra bang cach bat log '
                  '(Microsoft.EntityFrameworkCore.Database.Command) roi goi endpoint.')
        add(p._p)
        made += 1

    d.save(a.rds)
    print('Da sinh {} muc III.n vao {}'.format(made, a.rds))
    print('  man co anh mockup     : {}/{}'.format(made - no_shot, made))
    print('  man co bang field Zod : {}/{}'.format(made - no_field, made))


if __name__ == '__main__':
    main()
