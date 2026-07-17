"""Xoa MOI dau vet du an mau khoi RDS/SDS + dat lai tieu de cho GymMaster.

Template cua thay chua san vi du cua 2 du an khac:
  - GAMS (Global Access Management System): tieu de, man Setting List/Details
  - Cafeteria Ordering System: UC-5 Order a Meal, UC-6 Register for Payroll
    Deduction, actor "Patron"
Khong xoa la nop nham du an nguoi khac — 'Patron' xuat hien toi 114 lan.

Dung: python strip_samples.py <file.docx> --title "GymMaster" [--subtitle "..."]
"""
import argparse
import sys

try:
    import docx
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    sys.exit('Thieu python-docx. Chay: uv run --with python-docx python ' + __file__)

# Bat cu doan/bang nao chua mot trong nhung tu nay deu la cua du an mau.
SAMPLE = ('Patron', 'Cafeteria', 'Payroll', 'Order a Meal', 'GAMS',
          'Global Access Management', 'COS ', 'Setting List', 'Setting Details',
          'Prithvi Raj', 'Nancy Anderson', 'KienNTHE11', 'MinhNNT',
          'daily special', 'meal order', 'payroll deduction')

# Muc mau trong phan II cua template
SAMPLE_HEADINGS = ('2. Common Functions', '2.1 UC-2_Login System',
                   '3. Patron Feature', '3.1 UC-5_Order a Meal',
                   '3.2 UC-6_Register for Payroll Deduction',
                   '1. <<Feature Name>>', '1.1 <<UseCaseCode_UC Name>>')


def text_of(el):
    return ''.join(n.text or '' for n in el.iter(qn('w:t')))


def is_sample(el):
    t = text_of(el)
    return any(k in t for k in SAMPLE)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('path')
    ap.add_argument('--title', default='GymMaster')
    ap.add_argument('--subtitle', default='Gym Management Web System')
    a = ap.parse_args()

    d = docx.Document(a.path)
    body = d.element.body

    removed = 0
    for el in list(body.iterchildren()):
        if el.tag not in (qn('w:p'), qn('w:tbl')):
            continue
        t = text_of(el).strip()
        if any(t.startswith(h) for h in SAMPLE_HEADINGS) or is_sample(el):
            # Giu lai muc dinh nghia template (Functional Description Contents)
            # vi do la HUONG DAN, khong phai du lieu mau cua du an khac.
            if t.startswith('Use Case ID and Name') or t.startswith('Author and Date'):
                continue
            body.remove(el)
            removed += 1

    # Dat lai tieu de. Luu y: dong 'Global Access Management System (GAMS)' da bi
    # vong xoa o tren cuon di roi -> phai CHEN LAI ten du an, khong thi file
    # khong con ten du an nao.
    import copy
    from datetime import date

    title_p = next((p for p in d.paragraphs[:20] if p.style.name == 'Title'), None)
    has_name = any(a.title in p.text for p in d.paragraphs[:20])
    if title_p is not None and not has_name:
        new = copy.deepcopy(title_p._p)
        title_p._p.addnext(new)
        np = docx.text.paragraph.Paragraph(new, title_p._parent)
        for r in np.runs:
            r.text = ''
        if np.runs:
            np.runs[0].text = '{} — {}'.format(a.title, a.subtitle)
        else:
            np.add_run('{} — {}'.format(a.title, a.subtitle))
        print('Da chen lai ten du an (dong GAMS bi xoa cung voi mau)')

    for p in d.paragraphs[:20]:
        if p.text.strip().startswith('<<Project name>>'):
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = a.title
        # 'Hanoi, August 2022' la ngay cua template
        if 'Hanoi,' in p.text and '202' in p.text:
            txt = '– Hanoi, {} –'.format(date.today().strftime('%B %Y'))
            for r in p.runs[1:]:
                r.text = ''
            if p.runs:
                p.runs[0].text = txt

    d.save(a.path)

    # kiem chung
    d2 = docx.Document(a.path)
    full = '\n'.join(p.text for p in d2.paragraphs)
    for t in d2.tables:
        for r in t.rows:
            for c in r.cells:
                full += '\n' + c.text
    left = {k: full.count(k) for k in ('Patron', 'Cafeteria', 'Payroll', 'GAMS',
                                       'Teacher', 'Setting List') if full.count(k)}
    print('Da xoa {} phan tu mau khoi {}'.format(removed, a.path))
    if left:
        print('CON SOT:', left)
    else:
        print('Sach: khong con dau vet du an mau nao.')


if __name__ == '__main__':
    main()
