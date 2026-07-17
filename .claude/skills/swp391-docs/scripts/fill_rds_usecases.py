"""Dien RDS phan II (Requirement Specifications): moi use case 1 bang 15 dong
+ 1 bang Business Rules, sinh tu 03_SRS_USE_CASES.md + specs/00X/spec.md.

Nguon (khong bia):
  UC list + chi tiet   <- docs/init/03_SRS_USE_CASES.md
  Business Rules       <- specs/00X/spec.md muc "3. Functional Requirements (EARS)"
  Exceptions           <- specs/00X/spec.md muc "7. Error Handling"
  Pre/Postconditions   <- muc "8. Acceptance Criteria (Given-When-Then)"

Dung: python fill_rds_usecases.py <rds.docx>
"""
import copy
import re
import sys
from pathlib import Path

try:
    import docx
    from docx.shared import Pt
except ImportError:
    sys.exit('Thieu python-docx. Chay: uv run --with python-docx python ' + __file__)

UC_MD = Path('docs/init/03_SRS_USE_CASES.md')
SPECS = Path('specs')

# UC -> thu muc spec. Suy tu ten UC; cai nao khong khop thi bao ra.
SPEC_OF = {
    'login': '001-auth-rbac', 'logout': '001-auth-rbac', 'password': '001-auth-rbac',
    'user account': '002-member-management', 'staff account': '002-member-management',
    'member profile': '002-member-management', 'pt profile': '002-member-management',
    'membership package': '003-membership-billing', 'sell membership': '003-membership-billing',
    'renew membership': '003-membership-billing', 'revenue': '008-dashboard-audit',
    'check-in': '004-checkin',
    'assign pt': '005-pt-training', 'assigned members': '005-pt-training',
    'workout plan': '005-pt-training', 'trainer note': '005-pt-training',
    '360': '006-progress-tracking', 'progress': '006-progress-tracking',
    'calorie': '007-nutrition-calorie', 'meal log': '007-nutrition-calorie',
    'food item': '007-nutrition-calorie', 'custom food': '007-nutrition-calorie',
    'barcode': '007-nutrition-calorie',
    'audit log': '008-dashboard-audit', 'reminder': '008-dashboard-audit',
    'image food': '009-image-food-recognition',
    'online payment': '010-online-payment-vnpay', 'vnpay': '010-online-payment-vnpay',
    'cancel': '003-membership-billing', 'self-service': '002-member-management',
}


def demd(s):
    s = re.sub(r'\*\*(.+?)\*\*', r'\1', s or '')
    s = re.sub(r'\*(.+?)\*', r'\1', s)
    return re.sub(r'\s+', ' ', s.replace('`', '')).strip()


def spec_for(uc_name):
    n = uc_name.lower()
    for k, v in SPEC_OF.items():
        if k in n:
            return v
    return None


def section(text, num):
    m = re.search(r'^## {}\..*?$(.*?)(?=^## \d+\.|\Z)'.format(num), text, re.M | re.S)
    return m.group(1) if m else ''


def read_uc_table():
    t = UC_MD.read_text(encoding='utf-8')
    blk = t[t.index('# 2. Use Case Overview'):t.index('# 3.')]
    out = []
    for ln in blk.splitlines():
        m = re.match(r'^\|\s*(UC-[0-9A-Z]+)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|',
                     ln.strip())
        if m:
            out.append(dict(zip(('id', 'name', 'actor', 'prio'),
                                [demd(x) for x in m.groups()])))
    return out


def read_uc_details():
    """Cac UC co muc chi tiet '## UC-01 — Login' trong 03_SRS_USE_CASES.md."""
    t = UC_MD.read_text(encoding='utf-8')
    det = {}
    for m in re.finditer(r'^## (UC-[0-9A-Z]+)\s*[—-]\s*(.+?)$(.*?)(?=^## |\Z)',
                         t, re.M | re.S):
        uid, body = m.group(1), m.group(3)
        d = {}
        for k in ('Objective', 'Trigger', 'Pre-condition', 'Post-condition'):
            mm = re.search(r'\|\s*' + k + r'\s*\|\s*(.+?)\s*\|', body)
            if mm:
                d[k] = demd(mm.group(1))
        for k in ('Main Flow', 'Exception Flow', 'Acceptance Criteria'):
            mm = re.search(r'\*\*' + k + r':?\*\*\s*(.+?)(?=\n\*\*|\n##|\Z)', body, re.S)
            if mm:
                d[k] = demd(mm.group(1))
        det[uid] = d
    return det


def frs_of(spec_dir):
    p = SPECS / spec_dir / 'spec.md'
    if not p.exists():
        return []
    t = p.read_text(encoding='utf-8')
    out = []
    for m in re.finditer(r'\*\*(FR-[A-Z]+-\d+)\s*\(([^)]+)\)[:*]*\*\*\s*(.+?)(?=\n- \*\*FR-|\n##|\Z)',
                         section(t, 3), re.S):
        out.append((m.group(1), m.group(2), demd(m.group(3))))
    return out


STOP = {'the', 'and', 'for', 'view', 'manage', 'add', 'set', 'basic', 'system',
        'user', 'member', 'a', 'to', 'in', 'of'}


def relevant_frs(frs, uc_name):
    """FR cua spec ma UC nay thuoc ve — LIET KE DU, de nguoi doc tu cat.

    Da thu loc theo tu khoa trong ten UC va HONG: 'Login' chi khop FR-RBAC-03
    (mot FR chang lien quan), con FR-AUTH-02 (dung nghia login) bi loai vi noi
    dung viet "gui email + mat khau dung", khong co chu "login".
    FR nao that su ap dung cho UC nao la viec CAN NGUOI DOC, may khong quyet duoc.
    Liet ke thua thi nhin thay va cat duoc; liet ke sai thi khong ai biet.
    """
    return frs


def errors_of(spec_dir):
    p = SPECS / spec_dir / 'spec.md'
    if not p.exists():
        return []
    t = section(p.read_text(encoding='utf-8'), 7)
    out = []
    for ln in t.splitlines():
        ln = ln.strip()
        if ln.startswith('- IF') or ln.startswith('- WHEN'):
            out.append(demd(ln.lstrip('- ')))
    return out


def acs_of(spec_dir):
    p = SPECS / spec_dir / 'spec.md'
    if not p.exists():
        return []
    t = section(p.read_text(encoding='utf-8'), 8)
    return [demd(m.group(1)) for m in re.finditer(r'\*\*(AC-\d+):?\*\*\s*(.+)', t)] or \
           [demd(ln.lstrip('- [ ]')) for ln in t.splitlines() if 'AC-' in ln]


def uniq_cells(row):
    """O gop trong Word dung chung mot <w:tc>; row.cells lap lai no nhieu lan.
    Khong loc thi ghi o sau se de len o truoc (hoac xoa trang no)."""
    out, seen = [], set()
    for c in row.cells:
        k = id(c._tc)
        if k not in seen:
            seen.add(k)
            out.append(c)
    return out


def set_cell(cell, text):
    cell.text = ''
    p = cell.paragraphs[0]
    parts = str(text).split('\n')
    for i, line in enumerate(parts):
        r = p.add_run(line)
        r.font.size = Pt(9)
        if i < len(parts) - 1:
            r.add_break()


def main():
    path = sys.argv[1]
    d = docx.Document(path)
    ucs = read_uc_table()
    det = read_uc_details()

    tpl_uc = d.tables[8]        # bang mau 15 dong (rong)
    tpl_br = d.tables[9]        # bang mau Business Rules

    body = d.element.body
    anchor = tpl_br._tbl        # chen sau bang mau BR

    thin = []
    n_made = 0
    for i, uc in enumerate(ucs, 1):
        sd = spec_for(uc['name'])
        dd = det.get(uc['id'], {})
        all_frs = frs_of(sd) if sd else []
        frs = relevant_frs(all_frs, uc['name'])
        errs = errors_of(sd) if sd else []
        acs = acs_of(sd) if sd else []
        gaps = []
        if not dd.get('Main Flow'):
            gaps.append('Main Flow')
        if not frs:
            gaps.append('Business Rules')
        if gaps:
            thin.append('{:<7} {:<32} thieu: {}'.format(
                uc['id'], uc['name'][:31], ', '.join(gaps)))

        n = i
        flow = dd.get('Main Flow', '')
        # Nguon viet lien: "1. Nhap email. 2. He thong kiem tra. 3. Xac dinh role."
        # Tach theo moc "<so>. " chu khong tach theo dau cham — tach sai se chen
        # so rac ("1. 1 / 2. Nhap email / 3. 2 ...").
        steps = [re.sub(r'^\d+\.\s*', '', s).strip(' .')
                 for s in re.split(r'\s+(?=\d+\.\s)', flow) if s.strip()] if flow else []
        steps = [s for s in steps if s]
        if steps:
            normal = '{}.0 {}\n'.format(n, uc['name']) + '\n'.join(
                '{}. {}'.format(k, s) for k, s in enumerate(steps, 1))
        else:
            normal = ('{}.0 {}\n[CAN BO SUNG] 03_SRS_USE_CASES.md chua co Main Flow '
                      'chi tiet cho UC nay.'.format(n, uc['name']))

        exc = []
        for j, e in enumerate(errs, 1):
            exc.append('{}.0.E{} {}'.format(n, j, e))
        if dd.get('Exception Flow'):
            exc.insert(0, dd['Exception Flow'])

        pre = dd.get('Pre-condition', '')
        pre = 'PRE-1: {}'.format(pre) if pre else \
              'PRE-1: Nguoi dung da dang nhap va co quyen phu hop ({})'.format(uc['actor'])
        post = dd.get('Post-condition', '')
        post = 'POST-1: {}'.format(post) if post else ''

        rows = [
            ('UC ID and Name:', '{}_{}'.format(uc['id'], uc['name'])),
            ('Created By:', 'BanhMiChao', 'Date Created:', '2026-06-13'),
            ('Primary Actor:', uc['actor'], 'Secondary Actors:',
             'System' if 'System' not in uc['actor'] else 'None'),
            ('Trigger:', dd.get('Trigger', 'Actor chon chuc nang {} tren giao dien'.format(uc['name']))),
            ('Description:', dd.get('Objective', uc['name'])),
            ('Preconditions:', pre),
            ('Postconditions:', post),
            ('Normal Flow:', normal),
            ('Alternative Flows:', 'None'),
            ('Exceptions:', '\n'.join(exc) if exc else 'None'),
            ('Priority:', uc['prio']),
            ('Frequency of Use:', ''),
            ('Business Rules:', ('[CAN CAT BOT] Toan bo FR cua specs/{}: '.format(sd)
                                 + ', '.join(f[0] for f in frs)) if frs else 'None'),
            ('Other Information:', 'Spec chi tiet: specs/{}/spec.md'.format(sd) if sd else ''),
            ('Assumptions:', '\n'.join(acs[:3]) if acs else ''),
        ]

        t = copy.deepcopy(tpl_uc._tbl)
        anchor.addnext(t)
        anchor = t
        tb = docx.table.Table(t, d)
        for r, vals in enumerate(rows):
            # BAY O GOP: row.cells tra ve 4 phan tu nhung o gop dung CHUNG mot _tc.
            # Ghi cells[1] roi xoa cells[2:] chinh la tu xoa mat chu vua ghi.
            cells = uniq_cells(tb.rows[r])
            for j, v in enumerate(vals):
                if j < len(cells):
                    set_cell(cells[j], v)

        if frs:
            bt = copy.deepcopy(tpl_br._tbl)
            anchor.addnext(bt)
            anchor = bt
            btb = docx.table.Table(bt, d)
            while len(btb.rows) - 1 < len(frs):
                btb.add_row()
            while len(btb.rows) - 1 > len(frs):
                btb._tbl.remove(btb.rows[-1]._tr)
            for r, (fid, kind, text) in enumerate(frs, 1):
                c = uniq_cells(btb.rows[r])
                for j, v in enumerate((fid, kind, text)):
                    if j < len(c):
                        set_cell(c[j], v)

        p = d.add_paragraph()
        p._p.addprevious(copy.deepcopy(p._p))
        n_made += 1

    d.save(path)

    def say(s):
        # Console Windows la cp1252 -> ky tu nhu 'o' trong "Member 360o" lam crash.
        sys.stdout.write(s.encode('ascii', 'replace').decode('ascii') + '\n')

    say('Da sinh {} bang use case vao {}'.format(n_made, path))
    if thin:
        say('')
        say('CAN BO SUNG TAY - {} UC thieu du lieu nguon:'.format(len(thin)))
        for t_ in thin:
            say('  ' + t_)


if __name__ == '__main__':
    main()
