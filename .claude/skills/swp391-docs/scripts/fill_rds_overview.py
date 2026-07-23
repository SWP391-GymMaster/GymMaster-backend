"""Dien RDS phan I (Overview) tu spec + code that.

Bang trong Template2_RDS Document.docx (da xac minh bang cach dump):
  0  Record of Changes        1  I.1.1 Actors           2  I.1.2 Use Cases
  3  I.2.2 Screen Desc        4  I.2.3 Screen Auth      5  I.2.4 Non-UI Functions
  6  I.3.1 Table Desc         7  I.3.2 Package Desc

Dung: python fill_rds_overview.py <template.docx> <out.docx> <inventory.csv>
"""
import argparse
import csv
import re
import shutil
import sys
from datetime import date
from pathlib import Path

try:
    import docx
    from docx.shared import Pt
except ImportError:
    sys.exit('Thieu python-docx. Chay: uv run --with python-docx python ' + __file__)

UC_MD = Path('docs/01-SRS-Requirements/use-cases/srs-use-cases.md')
SCHEMA_MD = Path('docs/02-SDD-Architecture/database-design/database-schema.md')
FEATURES = Path('backend/GymMaster.API/Features')
API = Path('backend/GymMaster.API')

ROLES = ['Admin', 'Staff', 'PT', 'Member']

SQL_TYPES = {'BIGINT', 'INT', 'SMALLINT', 'TINYINT', 'NVARCHAR', 'VARCHAR', 'DATETIME2',
             'DATE', 'TIME', 'BIT', 'DECIMAL', 'FLOAT', 'UNIQUEIDENTIFIER', 'IDENTITY'}


def set_cell(cell, text):
    """Ghi text vao o. '\\n' phai thanh line-break that: add_run('a\\nb') trong
    python-docx cho ra 'ab' dinh lien, khong xuong dong."""
    cell.text = ''
    p = cell.paragraphs[0]
    for i, line in enumerate(str(text).split('\n')):
        r = p.add_run(line)
        r.font.size = Pt(10)
        if i < len(str(text).split('\n')) - 1:
            r.add_break()


def resize(table, n_rows):
    """Cho bang co dung n_rows dong du lieu (khong ke header)."""
    while len(table.rows) - 1 < n_rows:
        table.add_row()
    while len(table.rows) - 1 > n_rows:
        table._tbl.remove(table.rows[-1]._tr)


def fill(table, rows):
    resize(table, len(rows))
    for i, vals in enumerate(rows, 1):
        cells = table.rows[i].cells
        for j, v in enumerate(vals):
            if j < len(cells):
                set_cell(cells[j], v)


# ---------------------------------------------------------------- doc nguon
def read_actors():
    t = UC_MD.read_text(encoding='utf-8')
    blk = t[t.index('# 1. Actors'):t.index('# 2.')]
    out = []
    for ln in blk.splitlines():
        m = re.match(r'^\|\s*([A-Za-z]+)\s*\|\s*(.+?)\s*\|$', ln.strip())
        if m and m.group(1) != 'Actor' and not set(m.group(1)) <= set('-'):
            out.append((m.group(1), demd(m.group(2))))
    return out


def read_ucs():
    t = UC_MD.read_text(encoding='utf-8')
    blk = t[t.index('# 2. Use Case Overview'):t.index('# 3.')]
    out = []
    for ln in blk.splitlines():
        m = re.match(r'^\|\s*(UC-[0-9A-Z]+)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|', ln.strip())
        if m:
            g = [demd(x) for x in m.groups()]
            d = dict(zip(('id', 'name', 'actor', 'prio'), g))
            # UC da go khoi pham vi thi khong liet ke — RDS phai ta he thong THAT
            # da giao. Cot uu tien danh dau '~~Removed~~'.
            if 'removed' in d['prio'].lower():
                continue
            out.append(d)
    return out


def demd(s):
    """Go cu phap markdown: Word khong hieu ** va `, chung hien ra thanh rac."""
    s = re.sub(r'\*\*(.+?)\*\*', r'\1', s)      # **dam**
    s = re.sub(r'\*(.+?)\*', r'\1', s)          # *nghieng*
    s = s.replace('`', '')                      # `code`
    s = re.sub(r'\s+', ' ', s)
    return s.strip(' .·')


def _from_markdown_table(block):
    """`users` viet duoi dang BANG markdown, 22 bang con lai viet dang gach giua.

    Khong xu ly rieng thi ca bang markdown bi do nguyen xi vao o mo ta:
    '| Cot | Kieu | Rang buoc | Ghi chu | |--...'
    """
    cols, pk, fk = [], [], []
    for ln in block.splitlines():
        ln = ln.strip()
        if not ln.startswith('|') or set(ln) <= set('|- '):
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        if not cells or cells[0].lower() in ('cột', 'cot', 'column'):
            continue
        name = cells[0].strip('`')
        if not name:
            continue
        cols.append(name)
        rest = ' '.join(cells[1:])
        if re.search(r'\bPK\b', rest):
            pk.append(name)
        if re.search(r'\bFK\b|→', rest):
            fk.append(name)
    return cols, pk, fk


def read_tables():
    """Bang 2.x trong database-schema.md: '## 2.13 `check_ins`' + mo ta."""
    t = SCHEMA_MD.read_text(encoding='utf-8')
    out = []
    for m in re.finditer(r'^## 2\.\d+\s+`(\w+)`\s*\n(.+?)(?=\n## |\Z)', t, re.M | re.S):
        name, block = m.group(1), m.group(2)
        if block.lstrip().startswith('|'):
            cols, pk, fk = _from_markdown_table(block)
            d = 'Các cột: ' + ', '.join(cols)
        else:
            flat = demd(' '.join(block.split()))
            # Dang viet la 'Id BIGINT PK' -> regex (\w+)\s+PK chop trung 'BIGINT'.
            # Phai cho phep co kieu SQL o giua, va loai kieu SQL ra khoi ten cot.
            pk = [g[0] if g[0] not in SQL_TYPES else g[1]
                  for g in re.findall(r'(\w+)(?:\s+(\w+))?\s+PK\b', flat)]
            pk = [p for p in pk if p and p not in SQL_TYPES]
            # Khoa kep viet kieu '**PK kép (UserId, RoleId)**' — khong bat thi
            # user_roles / cac bang noi se hien ra khong co PK.
            comp = re.search(r'PK\s*(?:kép|kep|composite)?\s*\(([^)]+)\)', flat)
            if comp:
                pk = [c.strip() for c in comp.group(1).split(',') if c.strip()]
            fk = [m for m in re.findall(r'(\w+)\s+(?:\w+\s+)?(?:\(FK|FK)\b', flat)
                  if m not in SQL_TYPES]
            d = flat.replace(' · ', '; ')[:300]
        if pk:
            d += '\nPrimary keys: ' + ', '.join(dict.fromkeys(pk))
        if fk:
            d += '\nForeign keys: ' + ', '.join(dict.fromkeys(fk))
        out.append((name, d))
    return out


def uc_feature(uc):
    n = uc['name'].lower()
    for k, v in (('login', 'Authentication'), ('logout', 'Authentication'),
                 ('password', 'Authentication'), ('account', 'User Management'),
                 ('member profile', 'Member Management'), ('pt profile', 'Trainer Management'),
                 ('package', 'Membership & Billing'), ('membership', 'Membership & Billing'),
                 ('payment', 'Membership & Billing'), ('vnpay', 'Membership & Billing'),
                 ('renew', 'Membership & Billing'), ('check-in', 'Check-in'),
                 ('assign', 'PT Training'), ('workout', 'PT Training'),
                 ('trainer note', 'PT Training'), ('progress', 'Progress Tracking'),
                 ('360', 'Progress Tracking'), ('calorie', 'Nutrition'),
                 ('meal', 'Nutrition'), ('food', 'Nutrition'), ('barcode', 'Nutrition'),
                 ('dashboard', 'Dashboard & Audit'), ('audit', 'Dashboard & Audit'),
                 ('reminder', 'Dashboard & Audit')):
        if k in n:
            return v
    return 'General'


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('template')
    ap.add_argument('out')
    ap.add_argument('inventory')
    ap.add_argument('--author', default='BanhMiChao')
    a = ap.parse_args()

    shutil.copy(a.template, a.out)
    d = docx.Document(a.out)
    T = d.tables

    # --- 0. Record of Changes -------------------------------------------
    today = date.today().isoformat()
    fill(T[0], [['1.0', '2026-06-13', 'A', a.author,
                 'Khoi tao RDS: overall requirement + spec cho function Iteration 1'],
                ['1.1', '2026-06-27', 'M', a.author,
                 'Cap nhat cho Iteration 2 (nutrition, progress, Member 360, billing FE)'],
                ['1.2', '2026-07-11', 'M', a.author,
                 'Cap nhat cho Iteration 3 (account/profile, admin, member profile)'],
                ['1.3', today, 'M', a.author,
                 'Iteration 4: dong bo toan bo spec theo code that; backend refactor '
                 'sang feature-based; bo sung VNPay, Gemini food scan, OTP reset']])

    # --- 1. I.1.1 Actors -------------------------------------------------
    actors = read_actors()
    fill(T[1], [[str(i), n, desc] for i, (n, desc) in enumerate(actors, 1)])

    # --- 2. I.1.2 Use Cases ----------------------------------------------
    ucs = read_ucs()
    fill(T[2], [[u['id'], uc_feature(u), u['name'],
                 'Actor: {} | Priority: {}'.format(u['actor'], u['prio'])] for u in ucs])

    # --- 3. I.2.2 Screen Descriptions ------------------------------------
    with open(a.inventory, encoding='utf-8-sig') as fh:
        inv = list(csv.DictReader(fh))
    screens = [r for r in inv if r['kind'] == 'Screen']
    sys.path.insert(0, str(Path(__file__).parent))
    from fill_tracking import NAMES, NON_UI
    fill(T[3], [[str(i), r['feature'].capitalize(), NAMES.get(r['name'], r['name']),
                 'Route {} — {}'.format(r['name'], r['actor'])]
                for i, r in enumerate(screens, 1)])

    # --- 4. I.2.3 Screen Authorization -----------------------------------
    hdr = T[4].rows[0].cells
    set_cell(hdr[0], 'Screen')
    for j, role in enumerate(ROLES, 1):
        if j < len(hdr):
            set_cell(hdr[j], role)
    rows = []
    for r in screens:
        nm = NAMES.get(r['name'], r['name'])
        allowed = r['actor']
        row = [nm]
        for role in ROLES:
            row.append('X' if (allowed == role or allowed == 'Anonymous'
                               or role in allowed.split('/')) else '')
        rows.append(row)
    fill(T[4], rows)

    # --- 5. I.2.4 Non-UI Functions ---------------------------------------
    fill(T[5], [[str(i), f, n, desc] for i, (n, f, _ac, desc) in enumerate(NON_UI, 1)])

    # --- 6. I.3.1 Table Descriptions -------------------------------------
    tbls = read_tables()
    fill(T[6], [['{:02d}'.format(i), n, desc] for i, (n, desc) in enumerate(tbls, 1)])

    # --- 7. I.3.2 Code Packages ------------------------------------------
    pkgs = []
    for p in sorted(FEATURES.iterdir()):
        if p.is_dir():
            n = len(list(p.glob('*.cs')))
            pkgs.append(('Features.' + p.name,
                         'Feature {}: controller + service + interface + DTO ({} file). '
                         'Namespace GymMaster.API.Features.{}'.format(p.name, n, p.name)))
    for extra, desc in (
            ('Common', 'Kieu dung chung xuyen feature: ServiceResult<T>, ApiResponse<T>, '
                       'PagedResult<T>, ApiControllerBase, AppClock (GMT+7), PersonValidation'),
            ('Infrastructure', 'Ket noi he thong ngoai: Cloudinary (avatar), Gemini Vision '
                               '(food scan), EmailSender (SMTP OTP), VnPayLibrary (HMAC-SHA512)'),
            ('Entities', 'Entity EF Core (POCO thuan, cau hinh bang Fluent API)'),
            ('Data', 'GymMasterDbContext + DatabaseSeeder'),
            ('Options', 'Cau hinh strongly-typed: Jwt, VnPay, Gemini, Cloudinary, Email, '
                        'GoogleAuth, CheckIn')):
        p = API / extra
        if p.is_dir():
            pkgs.append((extra, desc + ' ({} file)'.format(len(list(p.glob('*.cs'))))))
    fill(T[7], [['{:02d}'.format(i), n, desc] for i, (n, desc) in enumerate(pkgs, 1)])

    d.save(a.out)
    print('Da dien RDS phan I -> {}'.format(a.out))
    print('  I.1.1 Actors            : {}'.format(len(actors)))
    print('  I.1.2 Use Cases         : {}'.format(len(ucs)))
    print('  I.2.2 Screen Desc       : {}'.format(len(screens)))
    print('  I.2.3 Screen Auth       : {} x {} role'.format(len(screens), len(ROLES)))
    print('  I.2.4 Non-UI Functions  : {}'.format(len(NON_UI)))
    print('  I.3.1 Table Desc        : {}'.format(len(tbls)))
    print('  I.3.2 Code Packages     : {}'.format(len(pkgs)))


if __name__ == '__main__':
    main()
