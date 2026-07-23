"""Dien SDS (Template3) tu code that.

  Bang 0  Record of Changes
  Bang 1  I.1 Code Packages
  Bang 2  I.2.b Table Description
  Bang 3/4  II.b Class Specifications (mau -> nhan ban cho tung class)

Moi feature mot muc II.n: Class Diagram + Class Specifications (method that
+ mo ta suy tu ten/kieu tra ve) + Sequence Diagram + Database Queries.

Dung: python fill_sds.py <template.docx> <out.docx>
"""
import argparse
import copy
import re
import shutil
import sys
from datetime import date
from pathlib import Path

try:
    import docx
    from docx.shared import Pt
except ImportError:
    sys.exit('Thieu python-docx. Chay: uv run --with python-docx python ' + __file__)

API = Path('backend/GymMaster.API')
FEATURES = API / 'Features'
SCHEMA_MD = Path('docs/02-SDD-Architecture/database-design/database-schema.md')

sys.path.insert(0, str(Path(__file__).parent))
# Tai su dung bo doc class cua skill swp391-diagrams thay vi viet lai
# (no da qua kiem chung: chi lay method public, xu ly interface + record + property).
sys.path.insert(0, str(Path(__file__).resolve().parents[2]
                       / 'swp391-diagrams' / 'scripts'))
from class_mermaid import scan          # noqa: E402


def demd(s):
    s = re.sub(r'\*\*(.+?)\*\*', r'\1', s or '')
    return re.sub(r'\s+', ' ', s.replace('`', '')).strip(' .·')


def set_cell(cell, text):
    cell.text = ''
    p = cell.paragraphs[0]
    parts = str(text).split('\n')
    for i, line in enumerate(parts):
        r = p.add_run(line)
        r.font.size = Pt(9)
        if i < len(parts) - 1:
            r.add_break()


def uniq_cells(row):
    out, seen = [], set()
    for c in row.cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            out.append(c)
    return out


def fit(table, n):
    while len(table.rows) - 1 < n:
        table.add_row()
    while len(table.rows) - 1 > n:
        table._tbl.remove(table.rows[-1]._tr)


def fill(table, rows):
    fit(table, len(rows))
    for i, vals in enumerate(rows, 1):
        cs = uniq_cells(table.rows[i])
        for j, v in enumerate(vals):
            if j < len(cs):
                set_cell(cs[j], v)


def method_desc(name, ret, nargs):
    """Mo ta method suy tu ten + kieu tra ve — khong bia hanh vi."""
    verb = ('Lay danh sach' if name.startswith('List') else
            'Tao moi' if name.startswith('Create') else
            'Cap nhat' if name.startswith('Update') else
            'Xoa' if name.startswith('Delete') else
            'Tim' if name.startswith(('Find', 'Get', 'Search')) else
            'Kiem tra' if name.startswith(('Validate', 'Ensure', 'Check')) else
            'Xu ly')
    ret_txt = ret.replace('~', '<').replace('<', '<')
    return '{}. Input: {} tham so. Output: {}'.format(verb, nargs, ret_txt)


def read_tables():
    t = SCHEMA_MD.read_text(encoding='utf-8')
    out = []
    for m in re.finditer(r'^## 2\.\d+\s+`(\w+)`\s*\n(.+?)(?=\n## |\Z)', t, re.M | re.S):
        blk = m.group(2)
        if blk.lstrip().startswith('|'):
            cols = []
            for ln in blk.splitlines():
                ln = ln.strip()
                if ln.startswith('|') and not set(ln) <= set('|- '):
                    c = [x.strip() for x in ln.strip('|').split('|')]
                    if c and c[0].lower() not in ('cột', 'cot', 'column') and c[0]:
                        cols.append(c[0].strip('`'))
            d = 'Cac cot: ' + ', '.join(cols)
        else:
            d = demd(' '.join(blk.split())).replace(' · ', '; ')[:280]
        out.append((m.group(1), d))
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('template')
    ap.add_argument('out')
    ap.add_argument('--author', default='BanhMiChao')
    a = ap.parse_args()

    shutil.copy(a.template, a.out)
    d = docx.Document(a.out)
    T = d.tables

    # --- Bang 0: Record of Changes
    fill(T[0], [['2026-06-13', 'A', a.author, 'Khoi tao SDS: package + database design'],
                ['2026-06-27', 'M', a.author, 'Them code design Iteration 2'],
                ['2026-07-11', 'M', a.author, 'Them code design Iteration 3'],
                [date.today().isoformat(), 'M', a.author,
                 'Iteration 4: refactor backend sang feature-based; cap nhat toan bo '
                 'package/class theo cau truc moi']])

    # --- Bang 1: Code Packages
    pkgs = []
    for p in sorted(FEATURES.iterdir()):
        if p.is_dir():
            n = len(list(p.glob('*.cs')))
            pkgs.append(('GymMaster.API.Features.' + p.name,
                         'Feature {}: controller + service + interface + DTO ({} file)'
                         .format(p.name, n)))
    for extra, desc in (
            ('Common', 'ServiceResult<T>, ApiResponse<T>, PagedResult<T>, '
                       'ApiControllerBase, AppClock (GMT+7), PersonValidation'),
            ('Infrastructure', 'Cloudinary (avatar), Gemini Vision (food scan), '
                               'EmailSender (SMTP OTP), VnPayLibrary (HMAC-SHA512)'),
            ('Entities', 'Entity EF Core (POCO thuan, cau hinh bang Fluent API)'),
            ('Data', 'GymMasterDbContext + DatabaseSeeder'),
            ('Options', 'Cau hinh strongly-typed cho Jwt/VnPay/Gemini/Cloudinary/Email')):
        pp = API / extra
        if pp.is_dir():
            pkgs.append(('GymMaster.API.' + extra,
                         desc + ' ({} file)'.format(len(list(pp.glob('*.cs'))))))
    fill(T[1], [['{:02d}'.format(i), n, desc] for i, (n, desc) in enumerate(pkgs, 1)])

    # --- Bang 2: Table Description
    tbls = read_tables()
    fill(T[2], [['{:02d}'.format(i), n, desc] for i, (n, desc) in enumerate(tbls, 1)])

    # --- II. Code Designs: moi feature mot muc
    tpl_m = T[3]
    anchor = T[4]._tbl
    made = 0
    for p in sorted(FEATURES.iterdir()):
        if not p.is_dir():
            continue
        files = [f for f in p.glob('*.cs')]
        types = scan(files)
        svc = {k: v for k, v in types.items()
               if k.endswith('Service') and not k.startswith('I')}
        if not svc:
            continue

        h = d.add_paragraph(style='Heading 2')
        h.add_run('II.{} Feature {}'.format(made + 1, p.name))
        anchor.addnext(h._p)
        anchor = h._p

        for tag, txt in (
                ('Heading 3', 'a. Class Diagram'),
                ('Normal', '[Chen anh] Sinh bang: python .claude/skills/swp391-diagrams/'
                           'scripts/class_mermaid.py class out/diagrams/cls_{}.mmd '
                           '--feature {} — roi dan vao draw.io (Arrange > Insert > '
                           'Advanced > Mermaid) va xuat PNG.'.format(p.name, p.name)),
                ('Heading 3', 'b. Class Specifications')):
            q = d.add_paragraph(style=tag if tag != 'Normal' else None)
            q.add_run(txt)
            anchor.addnext(q._p)
            anchor = q._p

        for cname, info in sorted(svc.items()):
            h = d.add_paragraph(style='Heading 4')
            h.add_run('{} Class'.format(cname))
            anchor.addnext(h._p)
            anchor = h._p

            t = copy.deepcopy(tpl_m._tbl)
            anchor.addnext(t)
            anchor = t
            tb = docx.table.Table(t, d)
            ms = [(m, n, r) for m, n, r in info['methods'] if n >= 0]
            rows = [['{:02d}'.format(k), m, method_desc(m, r, n)]
                    for k, (m, n, r) in enumerate(ms, 1)] or \
                   [['01', '(khong co method public)', '-']]
            fill(tb, rows)

        for tag, txt in (
                ('Heading 3', 'c. Sequence Diagram(s)'),
                ('Normal', '[Chen anh] Sinh bang: python .claude/skills/swp391-diagrams/'
                           'scripts/sequence_mermaid.py <Controller> <Action> '
                           'out/diagrams/seq.mmd'),
                ('Heading 3', 'd. Database Queries'),
                ('Normal', '[CAN BO SUNG] Backend dung EF Core LINQ, khong co SQL tho. '
                           'Lay SQL that EF sinh ra bang cach bat log '
                           'Microsoft.EntityFrameworkCore.Database.Command roi goi endpoint.')):
            q = d.add_paragraph(style=tag if tag != 'Normal' else None)
            q.add_run(txt)
            anchor.addnext(q._p)
            anchor = q._p
        made += 1

    d.save(a.out)
    print('Da dien SDS -> {}'.format(a.out))
    print('  I.1 Code Packages   : {}'.format(len(pkgs)))
    print('  I.2 Table Desc      : {}'.format(len(tbls)))
    print('  II  Code Designs    : {} feature'.format(made))


if __name__ == '__main__':
    main()
